import os
from docx import Document
from ..images import describe_image


def parse_docx(filepath, out_dir, config):
    doc = Document(filepath)
    name = os.path.basename(filepath)
    img_mode = config["image"]["mode"]
    img_counter = 0
    lines = []

    lines.append(f"\n{'='*80}")
    lines.append(f"FILE: {name}")
    lines.append(f"{'='*80}\n")

    if img_mode != "skip":
        img_dir = os.path.join(out_dir, "images", name)
        os.makedirs(img_dir, exist_ok=True)
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    data = rel.target_part.blob
                    ext = os.path.splitext(rel.target_ref)[1] or ".png"
                    img_counter += 1
                    rel_path = os.path.join("images", name, f"embedded_{img_counter}{ext}")
                    img_path = os.path.join(out_dir, rel_path)
                    with open(img_path, "wb") as f:
                        f.write(data)
                    desc = describe_image(img_path, config)
                    lines.append(f"  [Image: {rel_path}]")
                    if desc:
                        lines.append(f"  → {desc}")
                    lines.append("")
                except Exception:
                    pass

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "")
            prefix = "#" * max(1, int(level)) if level.isdigit() else "#"
            lines.append(f"\n{prefix} {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        lines.append("\n[TABLE]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append("[/TABLE]\n")

    return "\n".join(lines)
